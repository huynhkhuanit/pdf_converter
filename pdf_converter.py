import os
import re
import logging
import warnings
import requests
from bs4 import BeautifulSoup
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from typing import Optional, Union, List
from urllib.parse import urljoin
import certifi
import logging
import sys

# Đặt mã hóa console thành UTF-8
if sys.stdout.encoding != 'utf-8':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

# Cấu hình logging với UTF-8
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('pdf_conversion.log', encoding='utf-8'),  # UTF-8 cho file log
        logging.StreamHandler(sys.stdout)  # Sử dụng console UTF-8
    ]
)

# Cấu hình Tesseract OCR
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
warnings.filterwarnings("ignore")

class PDFConverter:
    def __init__(self, output_dir: str = "output"):
        """Khởi tạo PDFConverter với thư mục output được chỉ định."""
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def _extract_pdf_url(self, url: str) -> Optional[str]:
        """Trích xuất URL của file PDF từ trang web."""
        try:
            response = requests.get(url, verify=False)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")
            
            # Tìm kiếm PDF trong các thẻ phổ biến
            for tag, attr in [("object", "data"), ("embed", "src"), ("a", "href")]:
                tags = soup.find_all(tag)
                for t in tags:
                    link = t.get(attr)
                    if not link:
                        continue
                    link = link.strip()
                    
                    # Xử lý JavaScript links
                    if link.lower().startswith("javascript:"):
                        match = re.search(r"'(\/FileData\/[^']+\.pdf)'", link)
                        if match:
                            return urljoin(url, match.group(1))
                    # Xử lý direct PDF links
                    elif ".pdf" in link.lower():
                        return urljoin(url, link)
            return None
        except Exception as e:
            logging.error(f"Lỗi khi trích xuất URL PDF: {str(e)}")
            return None

    def _download_pdf(self, pdf_url: str) -> Optional[str]:
        """Tải file PDF về máy."""
        try:
            local_path = os.path.join(self.output_dir, "temp.pdf")
            response = requests.get(pdf_url, verify=False)
            response.raise_for_status()
            
            with open(local_path, "wb") as f:
                f.write(response.content)
            return local_path
        except Exception as e:
            logging.error(f"Lỗi khi tải PDF: {str(e)}")
            return None

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Trích xuất text từ file PDF, sử dụng OCR nếu cần."""
        try:
            # Thử trích xuất text trực tiếp
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            
            # Nếu không trích xuất được text hoặc text quá ngắn, sử dụng OCR
            if len(text.strip()) < 30:
                logging.info("Không trích xuất được text trực tiếp, đang sử dụng OCR...")
                pages = convert_from_path(pdf_path)
                text = ""
                for i, page_img in enumerate(pages):
                    logging.info(f"Đang OCR trang {i+1}...")
                    page_text = pytesseract.image_to_string(page_img, lang="vie")
                    text += f"--- Trang {i+1} ---\n{page_text}\n"
            else:
                logging.info("Đã trích xuất text thành công bằng pdfplumber")
            
            return text
        except Exception as e:
            logging.error(f"Lỗi khi trích xuất text: {str(e)}")
            return ""

    def convert_from_url(self, url: str, output_format: str = "txt") -> bool:
        """Chuyển đổi PDF từ URL sang text."""
        try:
            # Trích xuất và tải PDF
            pdf_url = self._extract_pdf_url(url)
            if not pdf_url:
                logging.error("Không tìm thấy file PDF trong URL")
                return False
            
            pdf_path = self._download_pdf(pdf_url)
            if not pdf_path:
                return False
            
            # Trích xuất text
            text = self._extract_text_from_pdf(pdf_path)
            if not text:
                return False
            
            # Lưu kết quả
            base_name = os.path.splitext(os.path.basename(pdf_url))[0]
            if output_format == "txt":
                output_path = os.path.join(self.output_dir, f"{base_name}.txt")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
            elif output_format == "docx":
                output_path = os.path.join(self.output_dir, f"{base_name}.docx")
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)
            
            # Xóa file tạm
            os.remove(pdf_path)
            logging.info(f"Đã chuyển đổi thành công: {output_path}")
            return True
            
        except Exception as e:
            logging.error(f"Lỗi khi chuyển đổi từ URL: {str(e)}")
            return False

    def convert_from_file(self, pdf_path: str, output_format: str = "txt") -> bool:
        """Chuyển đổi file PDF local sang text."""
        try:
            if not os.path.exists(pdf_path):
                logging.error(f"Không tìm thấy file: {pdf_path}")
                return False
            
            # Trích xuất text
            text = self._extract_text_from_pdf(pdf_path)
            if not text:
                return False
            
            # Lưu kết quả
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            if output_format == "txt":
                output_path = os.path.join(self.output_dir, f"{base_name}.txt")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
            elif output_format == "docx":
                output_path = os.path.join(self.output_dir, f"{base_name}.docx")
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)
            
            logging.info(f"Đã chuyển đổi thành công: {output_path}")
            return True
            
        except Exception as e:
            logging.error(f"Lỗi khi chuyển đổi từ file: {str(e)}")
            return False

    def batch_convert(self, pdf_paths: List[str], output_format: str = "txt") -> dict:
        """Chuyển đổi hàng loạt file PDF."""
        results = {"success": [], "failed": []}
        for pdf_path in pdf_paths:
            if self.convert_from_file(pdf_path, output_format):
                results["success"].append(pdf_path)
            else:
                results["failed"].append(pdf_path)
        return results

def main():
    converter = PDFConverter()
    
    while True:
        print("\n=== PDF to Text Converter ===")
        print("1. Chuyển đổi từ URL")
        print("2. Chuyển đổi từ file local")
        print("3. Chuyển đổi hàng loạt file")
        print("4. Thoát")
        
        choice = input("\nChọn chức năng (1-4): ")
        
        if choice == "1":
            url = input("Nhập URL: ")
            format_choice = input("Chọn định dạng output (txt/docx): ").lower()
            converter.convert_from_url(url, format_choice)
            
        elif choice == "2":
            file_path = input("Nhập đường dẫn file PDF: ")
            format_choice = input("Chọn định dạng output (txt/docx): ").lower()
            converter.convert_from_file(file_path, format_choice)
            
        elif choice == "3":
            folder_path = input("Nhập đường dẫn thư mục chứa PDF: ")
            format_choice = input("Chọn định dạng output (txt/docx): ").lower()
            
            pdf_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) 
                        if f.lower().endswith('.pdf')]
            
            results = converter.batch_convert(pdf_files, format_choice)
            print(f"\nKết quả chuyển đổi:")
            print(f"Thành công: {len(results['success'])} file")
            print(f"Thất bại: {len(results['failed'])} file")
            
        elif choice == "4":
            break
            
        else:
            print("Lựa chọn không hợp lệ!")

if __name__ == "__main__":
    main() 